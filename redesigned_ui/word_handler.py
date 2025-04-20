import os
import re
import traceback
from datetime import datetime
from docx import Document

from core.settings_handler import settings

class WordHandler:
    def __init__(self):
        self.document = None
        self.movie_table = None
        self.series_table = None
        
    def open_document(self):
        """Open the Word document"""
        try:
            word_doc_path = settings.get("WORD_DOC_PATH", "")
            if not word_doc_path or not os.path.exists(word_doc_path):
                print(f"Document not found at path: {word_doc_path}")
                return False
                
            self.document = Document(word_doc_path)
            
            # Get tables from document
            movie_table_index = settings.get_movie_table_index()
            series_table_index = settings.get_series_table_index()
            
            if len(self.document.tables) >= movie_table_index:
                self.movie_table = self.document.tables[movie_table_index - 1]
            else:
                print(f"Movie table not found at index {movie_table_index}")
                return False
                
            if len(self.document.tables) >= series_table_index:
                self.series_table = self.document.tables[series_table_index - 1]
            else:
                print(f"Series table not found at index {series_table_index}")
                return False
                
            return True
        except Exception as e:
            print(f"Error opening document: {e}")
            traceback.print_exc()
            return False
    
    def save_document(self):
        """Save the Word document"""
        try:
            if self.document:
                word_doc_path = settings.get("WORD_DOC_PATH", "")
                self.document.save(word_doc_path)
                return True
            return False
        except Exception as e:
            print(f"Error saving document: {e}")
            traceback.print_exc()
            return False
    
    def close_document(self):
        """Close the document by setting it to None"""
        self.document = None
        self.movie_table = None
        self.series_table = None
    
    def get_next_movie_number(self):
        """Get the next available movie number"""
        try:
            if not self.movie_table:
                return 1
                
            # Skip header row and count data rows
            num_rows = len(self.movie_table.rows) - 1
            return num_rows + 1
        except Exception as e:
            print(f"Error getting next movie number: {e}")
            return 1
    
    def get_next_series_number(self):
        """Get the next available series number"""
        try:
            if not self.series_table:
                return 1
                
            # Skip header row and count data rows
            num_rows = len(self.series_table.rows) - 1
            return num_rows + 1
        except Exception as e:
            print(f"Error getting next series number: {e}")
            return 1
            
    def add_movie(self, movie_data, keep_open=False):
        """
        Add a movie entry to the Word document
        
        Args:
            movie_data (dict): Dictionary containing movie information with keys:
                - title: Movie title
                - duration: Movie duration (e.g., "2h 30m")
                - genres: Genre string (e.g., "Action/Adventure")
                - watch_date: Date watched (datetime or string)
                - release_date: Release date (datetime or string)
                - user_rating: User rating (float)
                - imdb_rating: IMDb rating (string, e.g. "8.5/10")
                - rt_rating: Rotten Tomatoes rating (string, e.g. "90%")
                - rewatch: Boolean indicating if this is a rewatch
            keep_open (bool): Whether to keep the document open after saving
        
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            if not self.movie_table:
                print("Movie table not available")
                return False
                
            # Add a new row to the table
            row = self.movie_table.add_row()
            
            # Format the watch date
            watch_date_str = ""
            if movie_data.get("watch_date"):
                if isinstance(movie_data["watch_date"], datetime):
                    watch_date_str = movie_data["watch_date"].strftime("%d.%m.%Y")
                else:
                    watch_date_str = movie_data["watch_date"]
            
            # Format the release date
            release_date_str = ""
            if movie_data.get("release_date"):
                if isinstance(movie_data["release_date"], datetime):
                    release_date_str = movie_data["release_date"].strftime("%d.%m.%Y")
                else:
                    release_date_str = movie_data["release_date"]
            
            # Prepare user rating
            user_rating_str = ""
            if movie_data.get("user_rating") is not None:
                try:
                    rating_val = float(movie_data["user_rating"])
                    user_rating_str = f"{rating_val:.1f}/10"
                except (ValueError, TypeError):
                    user_rating_str = str(movie_data["user_rating"])
            
            # Prepare IMDb rating
            imdb_rating_str = movie_data.get("imdb_rating", "")
            if imdb_rating_str and not imdb_rating_str.endswith("/10"):
                imdb_rating_str = f"{imdb_rating_str}/10"
                
            # Prepare RT rating
            rt_rating_str = movie_data.get("rt_rating", "")
            if rt_rating_str and not rt_rating_str.endswith("%"):
                rt_rating_str = f"{rt_rating_str}%"
            
            # Add movie name with "Rewatch" indicator if needed
            movie_name = movie_data.get("title", "")
            if movie_data.get("rewatch", False):
                movie_name = f"{movie_name} (Rewatch)"
                
            # Set values in the appropriate columns
            next_num = self.get_next_movie_number()
            
            # Set cells
            movie_columns = settings.get_movie_columns()
            row.cells[movie_columns["NO"]].text = str(next_num)
            row.cells[movie_columns["NAME"]].text = movie_name
            row.cells[movie_columns["TIME_DURATION"]].text = movie_data.get("duration", "")
            row.cells[movie_columns["GENRE"]].text = movie_data.get("genres", "")
            row.cells[movie_columns["WATCH_DATE"]].text = watch_date_str
            row.cells[movie_columns["RELEASE_DATE"]].text = release_date_str
            row.cells[movie_columns["RATE"]].text = user_rating_str
            row.cells[movie_columns["IMDB_RATING"]].text = imdb_rating_str
            row.cells[movie_columns["RT_RATING"]].text = rt_rating_str
            
            # Save the document
            self.save_document()
            
            # Close the document if not keep_open
            if not keep_open:
                self.close_document()
                
            return True
        except Exception as e:
            print(f"Error adding movie: {e}")
            traceback.print_exc()
            return False
    
    def add_series(self, series_data, keep_open=False):
        """
        Add a series entry to the Word document
        
        Args:
            series_data (dict): Dictionary containing series information with keys:
                - title: Series title
                - season: Season number
                - episodes: Number of episodes
                - genres: Genre string (e.g., "Drama/Crime")
                - start_date: Date started watching (datetime or string)
                - finish_date: Date finished watching (datetime or string)
                - first_air_date: First air date (datetime or string)
                - user_rating: User rating (float)
                - imdb_rating: IMDb rating (string, e.g. "8.5/10")
                - rt_rating: Rotten Tomatoes rating (string, e.g. "90%")
                - finished: Boolean indicating if watched completely
            keep_open (bool): Whether to keep the document open after saving
        
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            if not self.series_table:
                print("Series table not available")
                return False
                
            # Add a new row to the table
            row = self.series_table.add_row()
            
            # Format dates
            start_date_str = ""
            if series_data.get("start_date"):
                if isinstance(series_data["start_date"], datetime):
                    start_date_str = series_data["start_date"].strftime("%d.%m.%Y")
                else:
                    start_date_str = series_data["start_date"]
            
            finish_date_str = ""
            if series_data.get("finish_date"):
                if isinstance(series_data["finish_date"], datetime):
                    finish_date_str = series_data["finish_date"].strftime("%d.%m.%Y")
                else:
                    finish_date_str = series_data["finish_date"]
            
            first_air_date_str = ""
            if series_data.get("first_air_date"):
                if isinstance(series_data["first_air_date"], datetime):
                    first_air_date_str = series_data["first_air_date"].strftime("%d.%m.%Y")
                else:
                    first_air_date_str = series_data["first_air_date"]
            
            # Prepare user rating
            user_rating_str = ""
            if series_data.get("user_rating") is not None:
                try:
                    rating_val = float(series_data["user_rating"])
                    user_rating_str = f"{rating_val:.1f}/10"
                except (ValueError, TypeError):
                    user_rating_str = str(series_data["user_rating"])
            
            # Prepare IMDb rating
            imdb_rating_str = series_data.get("imdb_rating", "")
            if imdb_rating_str and not imdb_rating_str.endswith("/10"):
                imdb_rating_str = f"{imdb_rating_str}/10"
                
            # Prepare RT rating
            rt_rating_str = series_data.get("rt_rating", "")
            if rt_rating_str and not rt_rating_str.endswith("%"):
                rt_rating_str = f"{rt_rating_str}%"
            
            # Set values in the appropriate columns
            next_num = self.get_next_series_number()
            
            # Set finished status
            finished_str = "Yes" if series_data.get("finished", False) else "No"
            
            # Set cells
            series_columns = settings.get_series_columns()
            row.cells[series_columns["NO"]].text = str(next_num)
            row.cells[series_columns["NAME"]].text = series_data.get("title", "")
            row.cells[series_columns["SEASON"]].text = str(series_data.get("season", ""))
            row.cells[series_columns["EPISODE"]].text = str(series_data.get("episodes", ""))
            row.cells[series_columns["GENRE"]].text = series_data.get("genres", "")
            row.cells[series_columns["STARTING_DATE"]].text = start_date_str
            row.cells[series_columns["FINISHING_DATE"]].text = finish_date_str
            row.cells[series_columns["FIRST_EPI_DATE"]].text = first_air_date_str
            row.cells[series_columns["RATE"]].text = user_rating_str
            row.cells[series_columns["IMDB_RATING"]].text = imdb_rating_str
            row.cells[series_columns["RT_RATING"]].text = rt_rating_str
            row.cells[series_columns["FINISHED"]].text = finished_str
            
            # Save the document
            self.save_document()
            
            # Close the document if not keep_open
            if not keep_open:
                self.close_document()
                
            return True
        except Exception as e:
            print(f"Error adding series: {e}")
            traceback.print_exc()
            return False 